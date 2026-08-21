import os
import re
import io
from typing import Dict, List, Any, Optional
import docx
from pypdf import PdfReader

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX, PDF, or TXT file bytes."""
    ext = os.path.splitext(filename.lower())[1]
    text = ""

    if ext == ".docx":
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs)
        except Exception as err:
            print(f"[ResumeParser] Error reading DOCX file: {err}")

    elif ext == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages.append(extracted.strip())
            text = "\n".join(pages)
        except Exception as err:
            print(f"[ResumeParser] Error reading PDF file: {err}")

    else:
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception as err:
            print(f"[ResumeParser] Error decoding text file: {err}")

    return text.strip()

# ── Section Segmentation ────────────────────────────────────────────────────────

SECTION_PATTERNS = {
    "PROJECTS": re.compile(r'^(?:key\s+|technical\s+|academic\s+|personal\s+|featured\s+|notable\s+|selected\s+)*projects(?:\s*&.*|\s*:.*)?$', re.IGNORECASE),
    "EXPERIENCE": re.compile(r'^(?:work\s+|professional\s+|relevant\s+|industry\s+|internship\s+)*(?:experience|employment|work\s+history|internships)(?:\s*:.*)?$', re.IGNORECASE),
    "SKILLS": re.compile(r'^(?:technical\s+|core\s+)*(?:skills|competencies|technologies|tools\s*&\s*technologies|languages|technical\s+stack)(?:\s*:.*)?$', re.IGNORECASE),
    "EDUCATION": re.compile(r'^(?:education|academic\s+background|qualifications|coursework)(?:\s*:.*)?$', re.IGNORECASE),
    "SUMMARY": re.compile(r'^(?:professional\s+|career\s+)*(?:summary|profile|about\s+me|objective)(?:\s*:.*)?$', re.IGNORECASE),
}

# Regex to detect tech stack / tools descriptor lines inside project entries
STACK_PREFIX_RE = re.compile(
    r'^(?:tech(?:nology)?\s+stack|technologies|tools(?:\s+used)?|stack|built\s+with|environment|languages|frameworks|libraries|frontend|backend|database|platform)\s*[:\-–—]\s*(.*)$',
    re.IGNORECASE
)

def segment_resume_sections(text: str) -> Dict[str, List[str]]:
    """Segment raw resume text into distinct semantic sections based on document structure."""
    raw_lines = [l.strip() for l in text.split('\n') if l.strip()]

    current_section = "HEADER"
    section_lines: Dict[str, List[str]] = {
        "HEADER": [],
        "SUMMARY": [],
        "SKILLS": [],
        "PROJECTS": [],
        "EXPERIENCE": [],
        "EDUCATION": [],
        "OTHER": []
    }

    for line in raw_lines:
        cleaned_line = re.sub(r'^[#*_\-•\s\d\.\)]+', '', line).strip()
        
        matched_section = None
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(cleaned_line):
                matched_section = sec_name
                break
        
        if matched_section:
            current_section = matched_section
            continue

        section_lines[current_section].append(line)

    return section_lines

# ── Skill & Domain Extraction ───────────────────────────────────────────────────

DOMAIN_TAXONOMY = {
    "dbms": ["dbms", "database", "sql", "mysql", "postgresql", "mongodb", "acid", "relational", "nosql", "redis", "timescaledb", "indexing", "sharding", "normalization", "transactions"],
    "os": ["os", "operating system", "deadlock", "process", "thread", "concurrency", "semaphore", "mutex", "virtual memory", "linux", "ipc", "scheduling", "paging", "kernel"],
    "networking": ["network", "networking", "tcp", "udp", "http", "https", "socket", "websocket", "webrtc", "dns", "ip", "osi", "tls", "ssl", "rest", "grpc"],
    "oop": ["oop", "oops", "object oriented", "inheritance", "polymorphism", "encapsulation", "class", "abstraction", "interface", "design patterns", "solid"],
    "ds": ["data structures", "algorithms", "dsa", "tree", "graph", "binary search", "hash table", "heap", "stack", "queue", "dynamic programming", "recursion", "sorting"],
    "fullstack": ["react", "node", "express", "fastapi", "full stack", "fullstack", "frontend", "backend", "web", "restful", "microservices", "next.js", "vite", "typescript", "javascript"],
    "cloud": ["aws", "docker", "kubernetes", "ci/cd", "cloud", "ec2", "s3", "gcp", "azure", "serverless", "devops"]
}

def detect_domain_tags_from_text(text: str) -> List[str]:
    """Detect core CS & engineering domain tags present in resume text."""
    lower_text = text.lower()
    tags = set()

    for domain, keywords in DOMAIN_TAXONOMY.items():
        if any(re.search(r'\b' + re.escape(kw) + r'\b', lower_text) for kw in keywords):
            tags.add(domain)
            if domain == "dbms":
                tags.update(["sql", "database"])
            elif domain == "os":
                tags.update(["operating-systems"])
            elif domain == "networking":
                tags.update(["networks"])
            elif domain == "oop":
                tags.update(["oops", "object-oriented"])
            elif domain == "ds":
                tags.update(["data-structures", "algorithms"])

    return list(tags)

ALL_SKILL_PATTERNS = [
    "Python", "JavaScript", "TypeScript", "React", "React.js", "Next.js", "Node.js", "Express", "Express.js",
    "FastAPI", "Django", "Flask", "Java", "Spring", "Spring Boot", "C++", "C#", "Go", "Golang", "Rust",
    "MongoDB", "PostgreSQL", "MySQL", "Redis", "TimescaleDB", "Cassandra", "SQLite", "GraphQL", "RESTful APIs",
    "Docker", "Kubernetes", "AWS", "EC2", "S3", "GCP", "Azure", "Git", "GitHub", "CI/CD", "Linux",
    "WebSockets", "WebRTC", "Web Audio API", "BullMQ", "Celery", "Kafka", "RabbitMQ", "PyTorch", "TensorFlow",
    "Tailwind CSS", "Tailwind", "CSS3", "HTML5", "Redux", "Redux Toolkit", "Monaco Editor", "Vite",
    "GSAP", "ScrollTrigger", "Three.js", "Socket.io",
    "Operating Systems", "Data Structures", "Algorithms", "DBMS", "OOP", "Object-Oriented Programming",
    "Computer Networks", "Microservices", "System Design", "Distributed Systems"
]

def extract_skills_from_text(text: str) -> List[str]:
    """Extract recognized programming languages, frameworks, and technical competencies."""
    lower_text = text.lower()
    matched = []
    for skill in ALL_SKILL_PATTERNS:
        pattern = r'(?<!\w)' + re.escape(skill.lower()) + r'(?!\w)'
        if re.search(pattern, lower_text):
            canonical = skill.replace(".js", "").strip()
            if canonical not in matched:
                matched.append(canonical)
    return matched

# ── Structured Entity Extractors ───────────────────────────────────────────────

def extract_projects_from_text(text: str) -> List[Dict[str, Any]]:
    """
    Extract structured projects strictly from the PROJECTS section.
    Maintains clean boundaries so that work experience or tech stack headers are never parsed as projects.
    """
    sections = segment_resume_sections(text)
    project_lines = sections.get("PROJECTS", [])

    if not project_lines:
        project_lines = sections.get("OTHER", []) + sections.get("HEADER", [])

    projects = []
    current_proj = None

    for line in project_lines:
        is_bullet = bool(re.match(r'^(?:[-•*–—+>]|\d+[\.\)]|[^\w\s])\s*', line))
        bullet_text = re.sub(r'^(?:[-•*–—+>]|\d+[\.\)]|[^\w\s])\s*', '', line).strip()

        if is_bullet and current_proj:
            if current_proj["description"]:
                current_proj["description"] += " " + bullet_text
            else:
                current_proj["description"] = bullet_text
            current_proj["highlights"].append(bullet_text)
            
            for s in extract_skills_from_text(bullet_text):
                if s not in current_proj["techStack"]:
                    current_proj["techStack"].append(s)
            continue

        # Check if line is a Stack / Tools descriptor line (e.g. "Stack: React.js, TypeScript, GSAP...")
        m_stack = STACK_PREFIX_RE.match(line)
        if m_stack:
            tech_str = m_stack.group(1)
            extracted = extract_skills_from_text(tech_str)
            if not extracted:
                extracted = [t.strip() for t in re.split(r'[,/|]', tech_str) if len(t.strip()) > 1]
            if current_proj:
                for t in extracted:
                    if t not in current_proj["techStack"]:
                        current_proj["techStack"].append(t)
            continue

        # Header line for project: e.g. "Project Title | React, Node.js, FastAPI"
        parts = [p.strip() for p in re.split(r'\||—|–', line) if p.strip()]
        candidate_title = parts[0] if parts else line

        # If candidate title starts with "Stack:" or "Tech Stack:" etc.
        m_cand_stack = STACK_PREFIX_RE.match(candidate_title)
        if m_cand_stack:
            if current_proj:
                for t in extract_skills_from_text(m_cand_stack.group(1)):
                    if t not in current_proj["techStack"]:
                        current_proj["techStack"].append(t)
            continue

        # Extract inline parens tech stack: "Real-Time Chat (React, WebSocket)"
        paren_match = re.search(r'\(([^)]+)\)', candidate_title)
        inline_tech = []
        if paren_match:
            inline_tech = extract_skills_from_text(paren_match.group(1))
            candidate_title = re.sub(r'\s*\([^)]*\)', '', candidate_title).strip()

        candidate_title = re.sub(r'^[#*_\-•\s\d\.\)]+', '', candidate_title).strip()
        if len(candidate_title) < 3 or len(candidate_title) > 75:
            continue

        # If candidate_title is purely a comma-separated list of recognized skills
        extracted_skills = extract_skills_from_text(candidate_title)
        if len(extracted_skills) >= 2 and (',' in candidate_title or '/' in candidate_title):
            if current_proj:
                for s in extracted_skills:
                    if s not in current_proj["techStack"]:
                        current_proj["techStack"].append(s)
            continue

        # Extract stack from pipe suffix
        pipe_tech = []
        if len(parts) > 1:
            pipe_tech = extract_skills_from_text(" ".join(parts[1:]))

        combined_stack = list(dict.fromkeys(inline_tech + pipe_tech + extract_skills_from_text(line)))

        if current_proj and (current_proj["description"] or current_proj["techStack"]):
            projects.append(current_proj)

        current_proj = {
            "title": candidate_title,
            "techStack": combined_stack,
            "description": "",
            "highlights": [],
            "role": "Developer"
        }

    if current_proj and (current_proj["description"] or current_proj["techStack"]):
        projects.append(current_proj)

    return projects[:4]

def extract_experience_from_text(text: str) -> List[Dict[str, Any]]:
    """Extract structured work experience entries strictly from the EXPERIENCE section."""
    sections = segment_resume_sections(text)
    exp_lines = sections.get("EXPERIENCE", [])

    experience = []
    current_exp = None

    for line in exp_lines:
        is_bullet = bool(re.match(r'^(?:[-•*–—+>]|\d+[\.\)]|[^\w\s])\s*', line))
        bullet_text = re.sub(r'^(?:[-•*–—+>]|\d+[\.\)]|[^\w\s])\s*', '', line).strip()

        if is_bullet and current_exp:
            current_exp["bullets"].append(bullet_text)
            continue

        parts = [p.strip() for p in re.split(r'\||—|–', line) if p.strip()]
        if parts:
            if current_exp and (current_exp["role"] or current_exp["bullets"]):
                experience.append(current_exp)
            current_exp = {
                "role": parts[0],
                "company": parts[1] if len(parts) > 1 else "",
                "bullets": []
            }

    if current_exp and (current_exp["role"] or current_exp["bullets"]):
        experience.append(current_exp)

    return experience

# ── High-Precision Question Synthesis Envelope ─────────────────────────────────

def synthesize_questions_from_resume(text: str, role: str = "Software Engineer", count: int = 5) -> Dict[str, Any]:
    """Synthesize candidate-specific interview questions grounded strictly in segmented resume structure."""
    if not text:
        text = role

    domain_tags = detect_domain_tags_from_text(text)
    found_skills = extract_skills_from_text(text)
    extracted_projects = extract_projects_from_text(text)
    extracted_experience = extract_experience_from_text(text)

    questions = []

    # 1. Project-Grounded Questions (Track 3)
    if extracted_projects:
        for proj in extracted_projects[:3]:
            stack_str = ", ".join(proj["techStack"]) if proj["techStack"] else "relevant technologies"
            questions.append({
                "questionText": f"In your project '{proj['title']}', walk me through the end-to-end architecture, data flow, and how you engineered the core features using {stack_str}.",
                "track": "project",
                "dimension": "architecture",
                "expectedConcepts": [f"Architecture of {proj['title']}", f"Usage of {stack_str}", "Performance optimization and system trade-offs"],
                "keywords": [s.lower() for s in proj["techStack"]] + ["architecture", "bottleneck", "data flow", "optimization"],
                "referenceAnswer": f"Candidate should describe hands-on architectural decisions, data flow, and trade-offs in {proj['title']}.",
                "projectContext": proj
            })

    # 2. Subject/Skill Walkthrough (Track 2)
    if found_skills:
        skill1 = found_skills[0]
        questions.append({
            "questionText": f"According to your resume, you have experience with {skill1}. Can you walk us through a key system or challenge where you implemented {skill1}?",
            "track": "subject",
            "expectedConcepts": [f"Deep practical familiarity with {skill1}", "Component design and workflow"],
            "keywords": [skill1.lower(), "architecture", "implementation", "design", "workflow"],
            "referenceAnswer": f"Candidate should demonstrate practical experience and sound design with {skill1}."
        })

    # 3. Behavioral / HR Context (Track 1)
    exp_context = f" at {extracted_experience[0]['company']}" if extracted_experience and extracted_experience[0].get("company") else ""
    questions.append({
        "questionText": f"To start off, please introduce yourself and walk me through your technical background, core engineering strengths, and what drives your passion for {role} positions.",
        "track": "hr",
        "expectedConcepts": ["Clear situational narrative", "Relevant technical stack", "Motivation & achievements"],
        "keywords": ["background", "strengths", "experience", "education", "passion", "skills"],
        "referenceAnswer": "Candidate should deliver a structured 60-90 second introduction covering education, core technical skills, recent projects/experience, and career goals."
    })

    return {
        "questions": questions[:count],
        "domainTags": domain_tags,
        "skills": found_skills,
        "projects": extracted_projects,
        "experience": extracted_experience,
        "summary": f"Identified {len(extracted_projects)} technical projects, {len(extracted_experience)} work experience roles, and {len(found_skills)} skills across {len(domain_tags)} domains."
    }
